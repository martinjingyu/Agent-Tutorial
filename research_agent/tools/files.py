from __future__ import annotations

import fnmatch
import json
import time
from pathlib import Path

from ..safety import resolve_workspace_path
from .registry import json_result, registry


_SKIP_DIRS = {".git", "__pycache__", "node_modules", ".venv", "venv"}
_MAX_SNIPPET_CHARS = 500
_SEARCH_REPEAT_STATE: dict[str, tuple[tuple[object, ...], int]] = {}


def _resolve_readable_path(raw_path: str | None) -> Path:
    if not raw_path:
        raise ValueError("path is required")
    try:
        return resolve_workspace_path(raw_path)
    except ValueError:
        return Path(raw_path).expanduser().resolve()


def _extract_docx_text(path: Path) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells]
                table_texts.append(" | ".join(row_texts))
        parts = []
        if paragraphs:
            parts.append("\n".join(paragraphs))
        if table_texts:
            parts.append("\n--- TABLES ---\n" + "\n".join(table_texts))
        return "\n".join(parts)
    except ImportError:
        return "[ERROR: python-docx not installed. Install with: pip install python-docx]"
    except Exception as exc:
        return f"[ERROR extracting .docx content: {exc}]"


def _extract_pdf_text(path: Path) -> str:
    try:
        import fitz
        doc = fitz.open(str(path))
        pages = [page.get_text().strip() for page in doc if page.get_text().strip()]
        doc.close()
        return "\n\n".join(pages)
    except ImportError:
        return "[ERROR: PyMuPDF not installed. Run: pip install pymupdf]"
    except Exception as exc:
        return f"[ERROR extracting PDF: {exc}]"


def _read_file(args: dict, runtime: dict) -> str:
    raw_path = args.get("path")
    max_chars = int(args.get("max_chars") or 20000)
    offset = max(0, int(args.get("offset") or 0))
    path = _resolve_readable_path(raw_path)
    if not path.is_file():
        return json_result(success=False, error=f"File not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        text = _extract_docx_text(path)
    elif suffix == ".pdf":
        text = _extract_pdf_text(path)
    else:
        text = path.read_text(encoding="utf-8", errors="replace")
    total_chars = len(text)
    page = text[offset:]
    truncated = len(page) > max_chars
    return json_result(
        success=True,
        path=str(path),
        content=page[:max_chars],
        offset=offset,
        total_chars=total_chars,
        truncated=truncated,
        next_offset=(offset + max_chars if truncated else None),
    )


def _write_file(args: dict, runtime: dict) -> str:
    path = resolve_workspace_path(args.get("path"))
    content = args.get("content")
    if content is None:
        return json_result(success=False, error="content is required")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(str(content), encoding="utf-8")
    return json_result(success=True, path=str(path), bytes=len(str(content).encode("utf-8")))


def _list_files(args: dict, runtime: dict) -> str:
    root = resolve_workspace_path(args.get("path") or ".")
    max_results = int(args.get("max_results") or 200)
    if root.is_file():
        return json_result(success=True, files=[str(root)])
    files: list[str] = []
    for item in root.rglob("*"):
        if ".git" in item.parts or "__pycache__" in item.parts:
            continue
        files.append(str(item.relative_to(resolve_workspace_path("."))))
        if len(files) >= max_results:
            break
    return json_result(success=True, root=str(root), files=files, truncated=len(files) >= max_results)


def _coerce_int(value: object, default: int, min_value: int = 0, max_value: int | None = None) -> int:
    try:
        parsed = int(value)  # type: ignore[arg-type]
    except (TypeError, ValueError):
        parsed = default
    parsed = max(min_value, parsed)
    if max_value is not None:
        parsed = min(max_value, parsed)
    return parsed


def _iter_search_files(root: Path, file_glob: str | None = None):
    if root.is_file():
        candidates = [root]
    else:
        candidates = root.rglob("*")
    for path in candidates:
        if not path.is_file():
            continue
        if any(part in _SKIP_DIRS for part in path.parts):
            continue
        if file_glob and not fnmatch.fnmatch(path.name, file_glob) and not fnmatch.fnmatch(str(path), file_glob):
            continue
        yield path


def _context_window(lines: list[str], index: int, context: int) -> list[dict[str, object]]:
    if context <= 0:
        return []
    start = max(0, index - context)
    end = min(len(lines), index + context + 1)
    return [
        {"line": i + 1, "text": lines[i][:_MAX_SNIPPET_CHARS]}
        for i in range(start, end)
        if i != index
    ]


def _check_repeated_search(key: tuple[object, ...], runtime: dict) -> tuple[bool, str | None]:
    task_key = str(runtime.get("task_id") or runtime.get("session_id") or "default")
    last_key, count = _SEARCH_REPEAT_STATE.get(task_key, ((), 0))
    count = count + 1 if last_key == key else 1
    _SEARCH_REPEAT_STATE[task_key] = (key, count)
    if count >= 4:
        return False, (
            f"BLOCKED: exact same search repeated {count} times. "
            "The result has not changed; use the existing result, narrow the pattern, "
            "or read_file near a returned line number."
        )
    if count >= 3:
        return True, (
            f"Warning: exact same search repeated {count} times. "
            "Avoid re-search loops; proceed with the returned locations."
        )
    return True, None


def _search_files(args: dict, runtime: dict) -> str:
    query = str(args.get("pattern") or args.get("query") or "")
    if not query:
        return json_result(success=False, error="pattern/query is required")
    target = str(args.get("target") or "content")
    output_mode = str(args.get("output_mode") or "content")
    file_glob = args.get("file_glob")
    file_glob = str(file_glob) if file_glob else None
    root = _resolve_readable_path(args.get("path") or ".")
    limit = _coerce_int(args.get("limit", args.get("max_results")), 50, 1, 500)
    offset = _coerce_int(args.get("offset"), 0, 0)
    context = _coerce_int(args.get("context"), 0, 0, 20)
    timeout = float(args.get("timeout_seconds") or 10.0)
    search_key = (query, target, str(root), file_glob or "", limit, offset, output_mode, context)
    ok, repeat_msg = _check_repeated_search(search_key, runtime)
    if not ok:
        return json_result(success=False, error=repeat_msg)
    if not root.exists():
        return json_result(success=False, error=f"Path not found: {root}")

    deadline = time.monotonic() + timeout
    timed_out = False

    if target == "files":
        files: list[Path] = []
        for path in _iter_search_files(root):
            if time.monotonic() > deadline:
                timed_out = True
                break
            if fnmatch.fnmatch(path.name, query) or fnmatch.fnmatch(str(path), query) or query.lower() in path.name.lower():
                files.append(path)
        files.sort(key=lambda p: p.stat().st_mtime if p.exists() else 0, reverse=True)
        page = files[offset:offset + limit]
        return json_result(
            success=True,
            target="files",
            files=[str(p) for p in page],
            total_count=len(files),
            truncated=len(files) > offset + limit or timed_out,
            next_offset=(offset + limit if len(files) > offset + limit else None),
            warning=repeat_msg,
            timed_out=timed_out,
        )

    counts: dict[str, int] = {}
    matches: list[dict[str, object]] = []
    total_count = 0
    for path in _iter_search_files(root, file_glob=file_glob):
        if time.monotonic() > deadline:
            timed_out = True
            break
        try:
            lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
            char_offset = 0
            for line_index, line in enumerate(lines):
                if query.lower() in line.lower():
                    total_count += 1
                    counts[str(path)] = counts.get(str(path), 0) + 1
                    if total_count > offset and len(matches) < limit:
                        item: dict[str, object] = {
                            "path": str(path),
                            "line": line_index + 1,
                            "char_offset": char_offset,
                            "text": line[:_MAX_SNIPPET_CHARS],
                        }
                        ctx = _context_window(lines, line_index, context)
                        if ctx:
                            item["context"] = ctx
                        matches.append(item)
                    if output_mode == "files_only":
                        break
                char_offset += len(line) + 1
        except OSError:
            continue

    if output_mode == "count":
        return json_result(
            success=True,
            output_mode="count",
            counts=counts,
            total_count=total_count,
            truncated=timed_out,
            warning=repeat_msg,
            timed_out=timed_out,
        )
    if output_mode == "files_only":
        files = list(counts.keys())
        page = files[offset:offset + limit]
        return json_result(
            success=True,
            output_mode="files_only",
            files=page,
            total_count=len(files),
            truncated=len(files) > offset + limit or timed_out,
            next_offset=(offset + limit if len(files) > offset + limit else None),
            warning=repeat_msg,
            timed_out=timed_out,
        )

    truncated = total_count > offset + limit or timed_out
    return json_result(
        success=True,
        pattern=query,
        matches=matches,
        total_count=total_count,
        truncated=truncated,
        next_offset=(offset + limit if truncated else None),
        hint=(f"Results truncated. Use offset={offset + limit}, narrow pattern, or set file_glob." if truncated else None),
        warning=repeat_msg,
        timed_out=timed_out,
    )


def _patch_file(args: dict, runtime: dict) -> str:
    path = resolve_workspace_path(args.get("path"))
    old_text = args.get("old_text")
    new_text = args.get("new_text")
    if not old_text:
        return json_result(success=False, error="old_text is required")
    if new_text is None:
        return json_result(success=False, error="new_text is required")
    text = path.read_text(encoding="utf-8", errors="replace")
    count = text.count(str(old_text))
    if count != 1 and not args.get("replace_all"):
        return json_result(success=False, error=f"Expected one match, found {count}")
    updated = text.replace(str(old_text), str(new_text), -1 if args.get("replace_all") else 1)
    path.write_text(updated, encoding="utf-8")
    return json_result(success=True, path=str(path), replacements=count if args.get("replace_all") else 1)


registry.register(
    "read_file",
    {
        "description": "Read a text file. Use offset/max_chars to read only the relevant chunk of large cache/result files. Search first with search_files when you need a specific section.",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "max_chars": {"type": "integer", "default": 20000},
                "offset": {"type": "integer", "default": 0, "description": "Character offset to start reading from."},
            },
            "required": ["path"],
        },
    },
    _read_file,
)
registry.register(
    "write_file",
    {
        "description": "Write a UTF-8 text file inside the Code workspace. Use this to save reports or update project files.",
        "parameters": {
            "type": "object",
            "properties": {"path": {"type": "string"}, "content": {"type": "string"}},
            "required": ["path", "content"],
        },
    },
    _write_file,
)
registry.register(
    "list_files",
    {
        "description": "List files under a Code workspace path.",
        "parameters": {
            "type": "object",
            "properties": {"path": {"type": "string", "default": "."}, "max_results": {"type": "integer", "default": 200}},
            "required": [],
        },
    },
    _list_files,
)
registry.register(
    "search_files",
    {
        "description": (
            "Agent-friendly file search. Use before read_file on large cache/results: "
            "search by pattern, get path+line snippets, then read_file only the needed chunk. "
            "Supports target='content' for text search and target='files' for filename search. "
            "Use output_mode='files_only' or 'count' to reduce tokens; use context only when needed."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "pattern": {"type": "string", "description": "Case-insensitive text pattern, or glob/name pattern when target='files'."},
                "query": {"type": "string", "description": "Backward-compatible alias for pattern."},
                "target": {"type": "string", "enum": ["content", "files"], "default": "content"},
                "path": {"type": "string", "default": "."},
                "file_glob": {"type": "string", "description": "Optional file filter for content search, e.g. '*.txt' or '*.py'."},
                "limit": {"type": "integer", "default": 50},
                "offset": {"type": "integer", "default": 0},
                "output_mode": {"type": "string", "enum": ["content", "files_only", "count"], "default": "content"},
                "context": {"type": "integer", "default": 0},
                "max_results": {"type": "integer", "default": 50, "description": "Backward-compatible alias for limit."},
                "timeout_seconds": {"type": "number", "default": 10.0},
            },
            "required": [],
        },
    },
    _search_files,
)
registry.register(
    "patch_file",
    {
        "description": "Replace text in a Code workspace file. Requires a unique old_text unless replace_all is true.",
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "old_text": {"type": "string"},
                "new_text": {"type": "string"},
                "replace_all": {"type": "boolean", "default": False},
            },
            "required": ["path", "old_text", "new_text"],
        },
    },
    _patch_file,
)


def _read_pdf(args: dict, runtime: dict) -> str:
    path = _resolve_readable_path(args.get("path"))
    max_chars = int(args.get("max_chars") or 50000)
    if not path.is_file():
        return json_result(success=False, error=f"File not found: {path}")
    if path.suffix.lower() != ".pdf":
        return json_result(success=False, error=f"Not a PDF file: {path}")
    text = _extract_pdf_text(path)
    pages = text.split("\n\n")
    truncated = len(text) > max_chars
    return json_result(
        success=True,
        path=str(path),
        pages=len(pages),
        content=text[:max_chars],
        truncated=truncated,
    )


def _read_url_pdf(args: dict, runtime: dict) -> str:
    import tempfile
    import urllib.request

    url = (args.get("url") or "").strip()
    if not url:
        return json_result(success=False, error="url is required")
    max_chars = int(args.get("max_chars") or 50000)

    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
    except Exception as exc:
        return json_result(success=False, error=f"Failed to download PDF: {exc}")

    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        try:
            text = _extract_pdf_text(tmp_path)
        finally:
            tmp_path.unlink(missing_ok=True)
    except Exception as exc:
        return json_result(success=False, error=f"Failed to extract PDF text: {exc}")

    truncated = len(text) > max_chars
    return json_result(
        success=True,
        url=url,
        pages=text.count("\n\n") + 1,
        content=text[:max_chars],
        truncated=truncated,
    )


registry.register(
    "read_pdf",
    {
        "description": (
            "Extract text from a PDF file using PyMuPDF. "
            "Returns full page text. Use for reading research papers, reports, or any PDF document."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Path to the PDF file"},
                "max_chars": {"type": "integer", "default": 50000},
            },
            "required": ["path"],
        },
    },
    _read_pdf,
)

registry.register(
    "read_url_pdf",
    {
        "description": (
            "Download and extract text from an online PDF given its URL. "
            "Useful for reading research papers, reports, or any PDF accessible via a direct link."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "Direct URL to the PDF file"},
                "max_chars": {"type": "integer", "default": 50000},
            },
            "required": ["url"],
        },
    },
    _read_url_pdf,
)

